#!/usr/bin/env python3
"""
Format Word documents with proper table borders and widths.

This script post-processes Pandoc-generated DOCX files to add:
- Table borders (single line, black)
- Proper table widths (6.5 inches to fit margins)
- Bold header rows
- Consistent cell padding

Usage:
    python format_docx.py input.docx output.docx

Requires:
    pip install python-docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def set_table_borders(table, border_size=4, border_color="000000"):
    """
    Add borders to all cells in a table.

    Args:
        table: docx Table object
        border_size: Border thickness in eighths of a point (4 = 0.5pt)
        border_color: Hex color code without #
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing borders
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)

    # Create new borders element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        border.set(qn('w:space'), '0')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def set_table_width(table, width_inches=6.5):
    """
    Set table to specified width.

    Args:
        table: docx Table object
        width_inches: Desired width in inches (6.5 fits 1" margins on letter)
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing width
    for child in list(tblPr):
        if child.tag == qn('w:tblW'):
            tblPr.remove(child)

    # Set new width (in twips: 1 inch = 1440 twips)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_table_alignment(table, alignment='center'):
    """
    Set table alignment on page.

    Args:
        table: docx Table object
        alignment: 'left', 'center', or 'right'
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing alignment
    for child in list(tblPr):
        if child.tag == qn('w:jc'):
            tblPr.remove(child)

    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    tblPr.append(jc)


def style_header_row(table, bold=True, background_color=None):
    """
    Style the header row of a table.

    Args:
        table: docx Table object
        bold: Whether to make header text bold
        background_color: Optional hex color for background (e.g., 'D9D9D9')
    """
    if not table.rows:
        return

    header_row = table.rows[0]

    for cell in header_row.cells:
        # Make text bold
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Set background color if specified
        if background_color:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), background_color)
            tcPr.append(shd)


def set_cell_padding(table, padding_pt=2):
    """
    Set cell padding for all cells.

    Args:
        table: docx Table object
        padding_pt: Padding in points
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing cell margins
    for child in list(tblPr):
        if child.tag == qn('w:tblCellMar'):
            tblPr.remove(child)

    # Set new margins (in twips: 1 point = 20 twips)
    tblCellMar = OxmlElement('w:tblCellMar')

    for margin in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{margin}')
        m.set(qn('w:w'), str(int(padding_pt * 20)))
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)

    tblPr.append(tblCellMar)


def auto_fit_columns(table):
    """
    Set columns to auto-fit content.

    Args:
        table: docx Table object
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set table layout to auto
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)


def format_academic_docx(input_path, output_path, options=None):
    """
    Apply academic formatting to a Word document.

    Args:
        input_path: Path to input DOCX file
        output_path: Path to output DOCX file
        options: Dictionary of formatting options
    """
    options = options or {}

    # Default options
    table_width = options.get('table_width', 6.5)
    border_size = options.get('border_size', 4)
    border_color = options.get('border_color', '000000')
    header_bold = options.get('header_bold', True)
    header_bg = options.get('header_background', None)  # e.g., 'D9D9D9' for light gray
    cell_padding = options.get('cell_padding', 2)
    center_tables = options.get('center_tables', True)

    # Load document
    doc = Document(input_path)

    # Process each table
    for i, table in enumerate(doc.tables):
        print(f"  Formatting table {i + 1}...")

        # Add borders
        set_table_borders(table, border_size, border_color)

        # Set width
        set_table_width(table, table_width)

        # Style header row
        style_header_row(table, header_bold, header_bg)

        # Set cell padding
        set_cell_padding(table, cell_padding)

        # Center table
        if center_tables:
            set_table_alignment(table, 'center')

    # Save document
    doc.save(output_path)
    print(f"Formatted document saved: {output_path}")
    print(f"  Tables processed: {len(doc.tables)}")


def main():
    """Main entry point."""
    if len(sys.argv) < 3:
        print("Usage: python format_docx.py input.docx output.docx [options]")
        print()
        print("Options (JSON format):")
        print('  --options \'{"table_width": 6.5, "border_size": 4}\'')
        print()
        print("Available options:")
        print("  table_width      : Width in inches (default: 6.5)")
        print("  border_size      : Border thickness (default: 4, = 0.5pt)")
        print("  border_color     : Hex color (default: '000000')")
        print("  header_bold      : Bold headers (default: true)")
        print("  header_background: Header bg color (default: none)")
        print("  cell_padding     : Padding in points (default: 2)")
        print("  center_tables    : Center tables (default: true)")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    # Parse options if provided
    options = {}
    if len(sys.argv) > 4 and sys.argv[3] == '--options':
        import json
        options = json.loads(sys.argv[4])

    # Verify input exists
    if not Path(input_path).exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    print(f"Formatting: {input_path}")
    format_academic_docx(input_path, output_path, options)


if __name__ == "__main__":
    main()
